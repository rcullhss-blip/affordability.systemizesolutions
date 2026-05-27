"""
Letter of Claim — First Legal Solicitors format.
Evidence-led, lender-type-aware LOC generator.
"""
import io
import os
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.analysis.lender_classifier import classify_lender, get_loc_argument, is_possible_intermediary

# First Legal Solicitors firm details
FIRM_NAME    = "First Legal Solicitors"
FIRM_ADDRESS = "8 Princes Parade, Liverpool, L3 1DL"
FIRM_LOGO    = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "..", "..", "first legal logo.png"
)

FIRM_FOOTER_LINE1 = (
    "Please ensure that all correspondence is sent to our Liverpool Office: "
    "8 Princes Parade, Liverpool, L3 1DL."
)
FIRM_FOOTER_LINE2 = (
    "First Legal Solicitors Ltd is authorised and regulated by the Solicitors Regulation Authority "
    "under registration number 634939. Registered in England and Wales with Company Number: "
    "10381298. Registered Office at 4th Floor, 8 Princes Parade, Liverpool, England, L3 1DL."
)

# ── Known lender addresses ─────────────────────────────────────────────────
LENDER_ADDRESSES = {
    "capital one":              "Trent House, Station Street, Nottingham, NG2 3HX",
    "vanquis":                  "Fairburn House, 5 Godwin Street, Bradford, BD1 2AH",
    "jaja":                     "27 Old Gloucester Street, Holborn, London, WC1N 3AX",
    "newday":                   "7 Handyside Street, London, N1C 4DA",
    "aqua":                     "7 Handyside Street, London, N1C 4DA",
    "zable":                    "Level 33, 25 Canada Square, Canary Wharf, London, E14 5LQ",
    "marbles":                  "7 Handyside Street, London, N1C 4DA",
    "fluid":                    "7 Handyside Street, London, N1C 4DA",
    "barclaycard":              "1 Churchill Place, London, E14 5HP",
    "barclays":                 "1 Churchill Place, London, E14 5HP",
    "lloyds":                   "25 Gresham Street, London, EC2V 7HN",
    "halifax":                  "Trinity Road, Halifax, HX1 2RG",
    "natwest":                  "250 Bishopsgate, London, EC2M 4AA",
    "santander":                "2 Triton Square, Regent's Place, London, NW1 3AN",
    "hsbc":                     "8 Canada Square, London, E14 5HQ",
    "tsb":                      "Henry Duncan House, 120 George Street, Edinburgh, EH2 4LH",
    "provident":                "No 1 Godwin Street, Bradford, BD1 2SU",
    "amigo":                    "Nova Building, 118–128 Commercial Road, Bournemouth, BH2 5LT",
    "very":                     "Skyways House, Speke Road, Speke, Liverpool, L70 1AB",
    "littlewoods":              "Skyways House, Speke Road, Speke, Liverpool, L70 1AB",
    "tesco bank":               "Interpoint, 22 Haymarket Yards, Edinburgh, EH12 5BH",
    "sainsbury":                "33 Holborn, London, EC1N 2HT",
    "advantage finance":        "Synergy House, Guildhall Road, Northampton, NN1 1DP",
    "premium credit":           "Premium Credit Limited, Ermyn Way, Leatherhead, Surrey, KT22 8UX",
    "moneybarn":                "Moneybarn House, Athena Drive, Tachbrook Park, Warwick, CV34 6RL",
    "motonovo":                 "One Central Square, Cardiff, CF10 1FS",
    "black horse":              "25 Gresham Street, London, EC2V 7HN",
    "secure trust bank":        "One Arleston Way, Shirley, Solihull, B90 4LH",
    "close motor finance":      "Close Brothers Ltd, 10 Crown Place, London, EC2A 4FT",
    "paragon motor finance":    "51 Homer Road, Solihull, West Midlands, B91 3QJ",
    "alphera financial":        "Summit ONE, Summit Avenue, Farnborough, GU14 0FB",
    "bmw financial services":   "Summit ONE, Summit Avenue, Farnborough, GU14 0FB",
    "mercedes benz financial":  "Tongwell, Milton Keynes, MK15 8BA",
    "mercedes benz fnancial":   "Tongwell, Milton Keynes, MK15 8BA",
    "toyota financial":         "Great Burgh, Burgh Heath, Epsom, Surrey, KT18 5UZ",
    "honda finance":            "Cain Road, Bracknell, Berkshire, RG12 1HL",
    "rci financial":            "Rivers Office Park, Denham Way, Maple Cross, Rickmansworth, WD3 9YS",
    "stellantis financial":     "Pinley House, 2 Sunbeam Way, Coventry, CV3 1ND",
    "fca automotive":           "Pinley House, 2 Sunbeam Way, Coventry, CV3 1ND",
    "fce bank":                 "Central Office, Arterial Road, Chelmsford, Essex, CM2 0RQ",
    "zopa":                     "1st Floor, One Southampton Row, London, WC1B 5HA",
    "oakbrook finance":         "Pacific House, Northgate, Aldridge, Walsall, WS9 8LT",
    "lendable":                 "67 Grosvenor Street, London, W1K 3JN",
    "118118 money":             "1 Canada Square, Canary Wharf, London, E14 5AB",
    "creation consumer":        "Chadwick House, Blenheim Court, Solihull, B91 2AA",
    "creation financial":       "Chadwick House, Blenheim Court, Solihull, B91 2AA",
    "ikano":                    "22 Birchin Lane, London, EC3V 9DU",
    "shawbrook":                "Lutea House, Warley Hill Business Park, Brentwood, Essex, CM13 3BE",
    "bamboo":                   "1 Hammersmith Broadway, London, W6 9DL",
    "tandem bank":              "1 Hammersmith Broadway, London, W6 9DL",
    "zempler bank":             "Skyways House, Speke Road, Speke, Liverpool, L70 1AB",
}

# ── Account type labels (for facilities table) ─────────────────────────────
_ACCOUNT_TYPE_LABELS = {
    "CREDIT_CARD":    "Credit Card",
    "PERSONAL_LOAN":  "Personal Loan",
    "HIRE_PURCHASE":  "Hire Purchase / Motor Finance",
    "MORTGAGE":       "Mortgage",
    "CURRENT_ACCOUNT":"Current Account",
    "PAYDAY_LOAN":    "High-Cost Short-Term Loan",
    "CATALOGUE":      "Catalogue / Mail Order",
    "MAIL_ORDER":     "Mail Order",
    "HOME_CREDIT":    "Home Credit",
    "OVERDRAFT":      "Overdraft",
    "TELECOM":        "Telecommunications",
    "UTILITY":        "Utility",
    "STORE_CARD":     "Store Card",
    "OTHER":          "Credit Account",
}

_FINANCIAL_TYPES_LOC = {
    "CREDIT_CARD", "PERSONAL_LOAN", "HIRE_PURCHASE", "PAYDAY_LOAN",
    "CATALOGUE", "MAIL_ORDER", "HOME_CREDIT", "OVERDRAFT", "STORE_CARD", "OTHER",
}

# ── Risk parameter sections (23.x) — included dynamically ─────────────────
RISK_PARAMS = {
    "23.1": {
        "flag_types": {"AP_MARKER", "ARRANGEMENT_TO_PAY", "DEBT_MANAGEMENT", "INSOLVENCY",
                       "ARRANGEMENT_TO_PAY_MARKER"},
        "title": "23.1 Personal Insolvency / Debt Management",
        "text": (
            "Our Client entered a reduced payment arrangement for existing credit commitments as a result "
            "of their financial circumstances, implying a risk of financial vulnerability which should have "
            "been an important factor in your affordability assessment. This marker demonstrates that our "
            "Client was unable to meet their existing obligations in full and was relying on reduced payment "
            "arrangements, which directly impacts their ability to service additional credit."
        ),
    },
    "23.2": {
        "flag_types": {"CCJ", "ACTIVE_CCJ", "MULTIPLE_CCJ", "MULTIPLE_CCJS", "ACTIVE_CCJ_ON_FILE",
                       "PUBLIC_RECORD_INSOLVENCY"},
        "title": "23.2 County Court Judgments and Public Records",
        "text": (
            "The credit report indicates that one or more adverse public records were registered against "
            "our Client prior to the Agreement. Any lender carrying out a reasonable Creditworthiness "
            "Assessment would have identified these public records as significant red flags indicating "
            "an inability to service existing debt obligations. The presence of such records should have "
            "been a clear contraindication to extending further credit, regardless of whether the entry "
            "constitutes a formal consumer CCJ or an insolvency-related public record."
        ),
    },
    "23.3": {
        "flag_types": {"DEFAULT_REGISTERED", "ACTIVE_DEFAULT", "DEFAULT",
                       "ACTIVE_ADVERSE_AT_LENDING", "ADVERSE_AT_LENDING"},
        "title": "23.3 Defaults and Adverse Credit Markers",
        "text": None,  # Generated dynamically — see _build_default_text()
    },
    "23.4": {
        "flag_types": {"REPEATED_MISSED_PAYMENTS", "ARREARS", "MISSED_PAYMENTS", "LATE_PAYMENTS",
                       "MISSED_PAYMENT", "REPEATED_MISSED_PAYMENT"},
        "title": "23.4 Arrears and Adverse Payment Conduct",
        "text": None,  # Generated dynamically — see _build_arrears_text()
    },
    "23.5": {
        "flag_types": {"GAMBLING", "GAMBLING_TRANSACTIONS"},
        "title": "23.5 Gambling",
        "text": (
            "Within our Client's transactional statements, there is evidence of frequent gambling. "
            "This constitutes a high-risk indicator for lending purposes and suggests erratic or "
            "impulsive financial behaviour which could jeopardise the ability to meet loan repayments. "
            "A lender conducting a thorough affordability assessment and reviewing transactional data "
            "should have identified this as a significant risk factor that ought to have resulted in the "
            "application being declined or subjected to enhanced scrutiny."
        ),
    },
    "23.6": {
        "flag_types": {"OVERDRAFT_USAGE", "OVERDRAFT", "PERSISTENT_OVERDRAFT"},
        "title": "23.6 Persistent Overdraft Usage",
        "text": (
            "Further review of our Client's transactional history shows frequent use of an overdraft "
            "facility, indicating ongoing cash-flow issues and an inability to manage finances within "
            "available means. Our Client's consistent use of an overdraft demonstrates that their "
            "expenditure regularly exceeded their income, leaving no financial buffer to service "
            "additional credit commitments. This should have been identified as a material affordability "
            "concern by any lender conducting a reasonable assessment."
        ),
    },
    "23.7": {
        "flag_types": {"PAYDAY_LOAN", "HIGH_COST_CREDIT", "REPEAT_PAYDAY"},
        "title": "23.7 High-Cost Short-Term Credit",
        "text": (
            "Our Client's credit file shows the use of high-cost short-term credit at or around the "
            "time of the Agreement. The use of payday-style lending is a well-recognised indicator of "
            "financial distress and an inability to manage day-to-day financial obligations from regular "
            "income. A lender exercising reasonable care would have identified this as a significant red "
            "flag and declined to extend further credit without conducting enhanced affordability checks."
        ),
    },
    "23.8": {
        "flag_types": {"DEBT_STACKING", "HIGH_UTILISATION", "RAPID_BORROWING",
                       "ELEVATED_UTILISATION", "REPEAT_BORROWING", "HIGH_CREDIT_UTILISATION",
                       "MULTIPLE_HARD_SEARCHES", "HARD_SEARCHES"},
        "title": "23.8 Concurrent Credit Commitments and Borrowing Pattern",
        "text": None,  # Generated dynamically — see _build_stacking_text()
    },
}


# ── Date / matching helpers ────────────────────────────────────────────────

def _parse_date_loc(val):
    if not val:
        return None
    if isinstance(val, date):
        return val
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%B %Y", "%b %Y"):
        try:
            return datetime.strptime(str(val), fmt).date()
        except ValueError:
            continue
    return None


def _fmt_date(val) -> str:
    if not val:
        return "—"
    if isinstance(val, (date, datetime)):
        return val.strftime("%d/%m/%Y")
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(str(val), fmt).strftime("%d/%m/%Y")
        except ValueError:
            pass
    return str(val)


def _lender_match_loc(a: str, b: str) -> bool:
    al, bl = a.lower().strip(), b.lower().strip()
    if al in bl or bl in al:
        return True
    words_a = [w for w in al.split() if len(w) > 3]
    return any(w in bl for w in words_a)


def _find_lender_account(accounts: list, lender_name: str) -> dict:
    """Fuzzy-match lender name to the best account record."""
    lender_lower = lender_name.lower().strip()
    # Exact match
    for acc in accounts:
        if acc.get("lender", "").lower().strip() == lender_lower:
            return acc
    # Substring match
    for acc in accounts:
        al = acc.get("lender", "").lower().strip()
        if lender_lower in al or al in lender_lower:
            return acc
    # Word-level match
    words = [w for w in lender_lower.split() if len(w) > 3]
    for acc in accounts:
        al = acc.get("lender", "").lower()
        if any(w in al for w in words):
            return acc
    return {}


def _get_lender_address(lender_name: str) -> str:
    name_lower = lender_name.lower()
    for key, addr in LENDER_ADDRESSES.items():
        if key in name_lower:
            return addr
    return ""


# ── Concurrent facilities helpers ─────────────────────────────────────────

def _concurrent_facilities(all_accounts: list, lending_date, exclude_lender: str) -> list:
    """Return financial accounts open at the time of the agreement, excluding the defendant lender."""
    result = []
    for acc in all_accounts:
        lender = acc.get("lender", "")
        if not lender:
            continue
        if _lender_match_loc(lender, exclude_lender):
            continue
        acc_type = (acc.get("account_type") or "OTHER").upper()
        if acc_type not in _FINANCIAL_TYPES_LOC:
            continue
        opened = _parse_date_loc(acc.get("opened_date"))
        # Skip accounts opened AFTER the agreement date
        if lending_date and opened and opened >= lending_date:
            continue
        result.append(acc)
    # Sort: defaults first (strongest adverse), then active, then others
    def _sort_key(a):
        s = (a.get("status") or "").upper()
        if s in ("DEFAULT", "DEFAULTED"):
            return 0
        ph = a.get("payment_history") or []
        missed = sum(1 for p in ph if str(p).upper() in {"D", "3", "4", "5", "6"})
        return 1 if missed >= 1 else 2
    result.sort(key=_sort_key)
    return result[:14]  # Cap for document readability


def _facility_status_notes(acc: dict) -> str:
    status = (acc.get("status") or "").upper()
    ph = acc.get("payment_history") or []
    missed = sum(1 for p in ph if str(p).upper() in {"D", "3", "4", "5", "6"})

    if status in ("DEFAULT", "DEFAULTED"):
        amount = acc.get("balance") or acc.get("arrears_amount")
        if amount and float(amount) > 0:
            return f"Defaulted — balance £{float(amount):,.0f}"
        return "Defaulted"
    if missed >= 3:
        return f"Active — {missed} missed payment(s) on record"
    if missed >= 1:
        return "Active — adverse payment marker(s) recorded"
    if status == "SETTLED":
        return "Settled"
    return "Active"


# ── Dynamic risk param text generators ───────────────────────────────────

def _build_default_text(matching_flags: list, cal: dict, opened_date_str: str) -> str:
    defaults_list = cal.get("active_defaults_list", [])
    n_defaults = len(defaults_list)

    if n_defaults > 0:
        detail_parts = []
        for d in defaults_list[:6]:
            lname = d.get("lender", "Unknown")
            ddate = _fmt_date(d.get("date"))
            amount = d.get("amount")
            part = f"{lname} (registered {ddate}"
            if amount and float(amount) > 0:
                part += f", balance outstanding £{float(amount):,.0f}"
            part += ")"
            detail_parts.append(part)
        s = "s" if n_defaults > 1 else ""
        details_str = "; ".join(detail_parts)
        return (
            f"Our Client's credit report evidences {n_defaults} default{s} registered prior to the "
            f"Agreement: {details_str}. "
            f"A default is one of the most serious adverse credit events that can appear on a consumer "
            f"credit file, demonstrating that our Client had already failed to meet their payment "
            f"obligations to an existing creditor to the point of formal account termination. The "
            f"presence of registered default{s} at the date of the Agreement is a clear and objective "
            f"indicator of pre-existing financial difficulty that any lender conducting a reasonable "
            f"Creditworthiness Assessment would have identified and acted upon. Extending further credit "
            f"in these circumstances was contrary to the requirements of FCA CONC 5.2A."
        )

    # Fall back to flag descriptions if no CAL data
    for f in matching_flags:
        if isinstance(f, dict) and f.get("description"):
            return (
                f"{f['description']}. "
                "This adverse indicator should have been identified during the Creditworthiness "
                "Assessment and should have prompted enhanced scrutiny or refusal of the application."
            )

    return (
        "Our Client's credit report shows adverse default markers registered at or around the "
        "time of the Agreement. A default is a serious adverse credit event indicating that our "
        "Client had already failed to meet their payment obligations with an existing creditor. "
        "Any lender conducting a reasonable Creditworthiness Assessment would have identified and "
        "acted upon these markers prior to approving the Agreement."
    )


def _build_arrears_text(matching_flags: list, cal: dict) -> str:
    missed = cal.get("missed_payments_6m_prior", 0)

    # Use the specific flag description if available (it contains lender-specific detail)
    for f in matching_flags:
        if isinstance(f, dict) and f.get("description"):
            base = f["description"]
            return (
                f"{base}. "
                "The presence of arrears and adverse payment conduct demonstrates that our Client "
                "was experiencing ongoing financial difficulty and was unable to meet existing "
                "payment obligations on time. A lender conducting a proper Creditworthiness "
                "Assessment would have identified this deteriorating payment pattern as a clear "
                "indicator of affordability concerns that ought to have been addressed — not "
                "disregarded — before approving further credit."
            )

    if missed > 0:
        return (
            f"Our Client's credit report shows {missed} adverse payment marker(s) recorded in the "
            f"period leading up to the Agreement. These arrears demonstrate that our Client was "
            f"experiencing financial difficulty and was unable to meet existing payment obligations. "
            "A lender conducting a proper Creditworthiness Assessment would have identified this as "
            "a clear indicator of affordability concerns."
        )

    return (
        "Our Client's credit report shows payment arrears at or around the time of the Agreement. "
        "These arrears demonstrate that our Client was experiencing ongoing financial difficulty "
        "and was unable to meet existing payment obligations on time. A lender conducting a proper "
        "Creditworthiness Assessment would have identified this as a clear affordability concern."
    )


def _build_stacking_text(matching_flags: list, cal: dict) -> str:
    account_count = cal.get("active_account_count", 0)
    searches = cal.get("hard_searches_90d_prior", 0)

    parts = []

    if account_count >= 3:
        parts.append(
            f"At the date of the Agreement, our Client held {account_count} concurrent credit "
            f"facilities with other lenders. The accumulation of multiple active credit commitments "
            f"is a recognised indicator of over-indebtedness and financial stress."
        )

    if searches >= 2:
        parts.append(
            f"In addition, {searches} credit application footprints were recorded in the 90 days "
            f"preceding the Agreement, indicating sustained and repeated credit-seeking behaviour "
            f"consistent with a client under significant financial pressure."
        )

    if not parts:
        # Use flag descriptions
        for f in matching_flags:
            if isinstance(f, dict) and f.get("description"):
                parts.append(f["description"])
                break

    if parts:
        body = " ".join(parts)
        return (
            f"{body} "
            "A lender exercising reasonable care and conducting a proper Creditworthiness Assessment "
            "would have recognised these patterns of concurrent borrowing and repeated credit-seeking "
            "behaviour as indicators that our Client was already financially over-committed at the "
            "time of the Agreement."
        )

    return (
        "At the time of the Agreement, our Client was maintaining multiple active credit facilities "
        "and exhibiting patterns of repeat credit-seeking behaviour consistent with ongoing financial "
        "difficulty. A lender exercising reasonable care would have identified this as a significant "
        "indicator of over-indebtedness and would have declined to extend further credit."
    )


def _get_risk_param_text(key: str, matching_flags: list, cal: dict, opened_date_str: str) -> str:
    """Return the appropriate text for a risk parameter section — dynamic or static."""
    if key == "23.3":
        return _build_default_text(matching_flags, cal, opened_date_str)
    if key == "23.4":
        return _build_arrears_text(matching_flags, cal)
    if key == "23.8":
        return _build_stacking_text(matching_flags, cal)
    # Static text for all other sections
    return RISK_PARAMS[key]["text"]


def _active_risk_sections(flags: list, cal: dict, opened_date_str: str) -> list:
    """Return ordered list of (key, title, text) tuples for active risk params."""
    flag_types_present: dict[str, list] = {}
    for f in flags:
        ftype = (f.get("type", "") if isinstance(f, dict) else str(f)).upper()
        flag_types_present.setdefault(ftype, []).append(f)

    sections = []
    for key, param in RISK_PARAMS.items():
        matching_types = set(flag_types_present.keys()) & param["flag_types"]
        if not matching_types:
            continue
        matching_flags = []
        for t in matching_types:
            matching_flags.extend(flag_types_present[t])
        text = _get_risk_param_text(key, matching_flags, cal, opened_date_str)
        sections.append((key, param["title"], text))
    return sections


# ── Document formatting helpers ───────────────────────────────────────────

def _set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15
    rpr = style.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:cs"),    "Arial")
    rpr.insert(0, rfonts)


def _para(doc, text="", bold=False, italic=False, size=11, space_before=0,
          space_after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold      = bold
        run.italic    = italic
        run.font.name = "Arial"
        run.font.size = Pt(size)
    return p


def _heading(doc, text: str, size=11, center=False, underline=True,
             space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold      = True
    run.underline = underline
    run.font.name = "Arial"
    run.font.size = Pt(size)
    return p


def _main_heading(doc, text: str):
    return _heading(doc, text, size=14, center=True, underline=True,
                    space_before=16, space_after=10)


def _numbered_item(doc, number: str, text: str, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after        = Pt(5)
    p.paragraph_format.space_before       = Pt(0)
    p.paragraph_format.line_spacing       = 1.15
    if indent:
        p.paragraph_format.left_indent       = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
    run_num = p.add_run(f"{number}\t")
    run_num.font.name = "Arial"
    run_num.font.size = Pt(11)
    run_text = p.add_run(text)
    run_text.font.name = "Arial"
    run_text.font.size = Pt(11)
    return p


def _borderless_table(doc, rows, col_widths=(Cm(5.0), Cm(11.0))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Normal Table"
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
        for cell, txt, bold in [(cells[0], label, True), (cells[1], value, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(txt)
            r.bold = bold
            r.font.name = "Arial"
            r.font.size = Pt(11)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "none")
                b.set(qn("w:sz"),    "0")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                tcBorders.append(b)
            tcPr.append(tcBorders)
    return table


def _facilities_table(doc, facilities: list):
    """Three-column evidence table: Lender | Account Type | Status / Notes."""
    headers = ("Lender / Creditor", "Account Type", "Status / Adverse Notes")
    rows_data = [
        (
            acc.get("lender", "—"),
            _ACCOUNT_TYPE_LABELS.get((acc.get("account_type") or "OTHER").upper(), "Credit Account"),
            _facility_status_notes(acc),
        )
        for acc in facilities
    ]
    all_rows = [headers] + rows_data
    col_widths = (Cm(6.0), Cm(4.5), Cm(5.5))

    table = doc.add_table(rows=len(all_rows), cols=3)
    table.style = "Normal Table"

    for i, (c1, c2, c3) in enumerate(all_rows):
        cells = table.rows[i].cells
        for j, (cell, val, w) in enumerate(zip(cells, [c1, c2, c3], col_widths)):
            cell.width = w
            p = cell.paragraphs[0]
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.bold = (i == 0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                if i == 0 and side == "bottom":
                    border.set(qn("w:val"),   "single")
                    border.set(qn("w:sz"),    "4")
                    border.set(qn("w:space"), "1")
                    border.set(qn("w:color"), "CBD5E1")
                else:
                    border.set(qn("w:val"),   "none")
                    border.set(qn("w:sz"),    "0")
                    border.set(qn("w:color"), "auto")
                tcBorders.append(border)
            tcPr.append(tcBorders)
    _para(doc, "", space_after=6)


def _horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_footer(doc):
    for section in doc.sections:
        section.footer_distance = Cm(1.0)
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            para.clear()
        p1 = footer.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after  = Pt(2)
        p1.paragraph_format.line_spacing = 1.0
        r1 = p1.add_run(FIRM_FOOTER_LINE1)
        r1.font.name = "Arial"
        r1.font.size = Pt(7)
        r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        p2 = footer.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        r2 = p2.add_run(FIRM_FOOTER_LINE2)
        r2.font.name = "Arial"
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        pPr = p1._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"),   "single")
        top.set(qn("w:sz"),    "4")
        top.set(qn("w:space"), "8")
        top.set(qn("w:color"), "CBD5E1")
        pBdr.append(top)
        pPr.append(pBdr)


# ── Evidence strength assessment ──────────────────────────────────────────

_CRITICAL_FLAG_TYPES = {
    "ACTIVE_ADVERSE_AT_LENDING", "ACTIVE_CCJ", "MULTIPLE_CCJS",
    "PUBLIC_RECORD_INSOLVENCY", "DEFAULT_REGISTERED", "REPEATED_MISSED_PAYMENTS",
}


def _has_critical_evidence(flags: list) -> bool:
    return any(
        (f.get("type", "") if isinstance(f, dict) else str(f)).upper() in _CRITICAL_FLAG_TYPES
        for f in flags
    )


# ── Main generator ─────────────────────────────────────────────────────────

def _estimate_agreement_date(acc: dict) -> str:
    """When open_date is absent, derive an approximate year from other account dates."""
    for field in ("default_date", "settled_date", "last_payment_date", "last_updated"):
        val = _parse_date_loc(acc.get(field))
        if val:
            # Defaults typically occur 1-3 yrs after origination; use the year as rough reference
            return f"on or around {val.year}"
    return "on or around the date of the Agreement (exact date to be confirmed)"


def _review_notice(doc, warnings: list[str]):
    """Insert a prominent internal-review banner before the letter body."""
    banner = doc.add_paragraph()
    banner.paragraph_format.space_before = Pt(0)
    banner.paragraph_format.space_after = Pt(2)
    pPr = banner._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFE5E5")
    pPr.append(shd)
    run = banner.add_run("⚠  INTERNAL — SOLICITOR REVIEW REQUIRED BEFORE SENDING")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    for w in warnings:
        wp = doc.add_paragraph()
        wp.paragraph_format.space_before = Pt(0)
        wp.paragraph_format.space_after = Pt(1)
        pPr2 = wp._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"), "FFE5E5")
        pPr2.append(shd2)
        wr = wp.add_run(f"•  {w}")
        wr.font.size = Pt(8)
        wr.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def generate_loc_docx(schema: dict, lender_result, review_warnings: list[str] | None = None) -> bytes:
    lender = (lender_result.lender_name if hasattr(lender_result, "lender_name")
              else lender_result.get("lender_name", "Unknown Lender"))
    flags  = (lender_result.risk_flags if hasattr(lender_result, "risk_flags")
              else lender_result.get("risk_flags", [])) or []

    client         = schema.get("client", {})
    client_name    = client.get("name") or "Our Client"
    client_address = client.get("address") or "—"
    matter_ref     = client.get("matter_ref") or "—"
    today_str      = date.today().strftime("%d %B %Y")

    all_accounts = schema.get("accounts", [])
    lender_acc   = _find_lender_account(all_accounts, lender)
    opened_date_raw = lender_acc.get("opened_date")
    opened_date_obj = _parse_date_loc(opened_date_raw)
    # Use formatted date or fall back to an estimated period if date is missing
    opened_date  = _fmt_date(opened_date_raw) if opened_date_raw else _estimate_agreement_date(lender_acc)
    acc_ref      = lender_acc.get("account_ref") or "—"

    # Computed-at-lending snapshot (attached during analysis)
    cal = lender_acc.get("computed_at_lending") or {}

    lender_address = _get_lender_address(lender)
    lender_type    = classify_lender(lender)

    # Evidence strength — drives para 11 wording and overall claim framing
    critical_evidence = _has_critical_evidence(flags)

    # Risk parameter sections — dynamic text generation
    active_sections = _active_risk_sections(flags, cal, opened_date)

    # Concurrent facilities for para 20
    concurrent = _concurrent_facilities(all_accounts, opened_date_obj, lender)

    doc = Document()
    _set_doc_defaults(doc)
    for section in doc.sections:
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.54)

    # ── Letterhead — logo left, date right ───────────────────────────────
    logo_path = os.path.normpath(FIRM_LOGO)
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = "Normal Table"
    logo_cell = hdr_table.rows[0].cells[0]
    logo_cell.width = Cm(8)
    if os.path.exists(logo_path):
        lp = logo_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run()
        lr.add_picture(logo_path, width=Cm(3.8))
    for cell in hdr_table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
    date_cell = hdr_table.rows[0].cells[1]
    date_cell.width = Cm(8)
    dp = date_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.paragraph_format.space_after = Pt(0)
    dr = dp.add_run(today_str)
    dr.font.name = "Arial"
    dr.font.size = Pt(11)

    _horizontal_rule(doc)

    # ── Lender address block ──────────────────────────────────────────────
    _para(doc, lender, bold=True, size=11, space_after=1)
    if lender_address:
        for line in lender_address.split(","):
            line = line.strip()
            if line:
                _para(doc, line, size=11, space_after=1)
    _para(doc, "", space_after=6)

    # ── Our Ref ───────────────────────────────────────────────────────────
    _para(doc, f"Our Ref:\t{matter_ref}", size=11, space_after=10)

    # ── Salutation ────────────────────────────────────────────────────────
    _para(doc, "Dear Sirs,", size=11, space_after=6)

    # ── Re: line ──────────────────────────────────────────────────────────
    _para(doc,
          f"Re: {client_name} – Claim Arising from Unaffordable Credit Agreement with {lender}",
          bold=True, size=11, space_after=10)

    # ── Client details ────────────────────────────────────────────────────
    _borderless_table(doc, [
        ("Our Client:",       client_name),
        ("Address:",          client_address),
        ("Agreement Number:", acc_ref),
    ])
    _para(doc, "", space_after=4)

    # ── LETTER OF CLAIM ───────────────────────────────────────────────────
    _main_heading(doc, "LETTER OF CLAIM")

    # Paragraphs 1–3: Introduction
    _numbered_item(doc, "1.",
        f"We act on behalf of our client, {client_name}, (\"Client\") in relation to a claim for "
        f"damages arising from the Client entering into a Consumer Credit Agreement (\"the Agreement\") "
        f"with you (\"Lender\").",
        indent=True)

    _numbered_item(doc, "2.",
        f"Our Client secured funding through you (\"Lender\") pursuant to the Agreement dated on or "
        f"around {opened_date}, which was regulated by the Consumer Credit Act 1974 (\"CCA 1974\").",
        indent=True)

    _numbered_item(doc, "3.",
        "Our Client believes that the Agreement was unaffordable at the time it was entered into "
        "and that you failed to carry out the creditworthiness and affordability assessment required "
        "of you by the Financial Conduct Authority.",
        indent=True)

    # ── Lien for Payment ──────────────────────────────────────────────────
    _heading(doc, "Lien for Payment of Fees", underline=True)

    _numbered_item(doc, "4.",
        "We hereby give notice that we have been engaged by our Client under a no-win, no-fee "
        "agreement (\"Retainer\"). Pursuant to the terms of our Retainer, any damages recovered "
        "on behalf of our Client must be paid to us directly, and not to our Client.",
        indent=True)

    _numbered_item(doc, "5.",
        "It follows that if you make payment of damages directly to our Client, we will enforce "
        "the lien against you in order to recover payment of our costs, which will therefore be "
        "payable in addition to any sum paid to our Client.",
        indent=True)

    _numbered_item(doc, "6.",
        "The firm's client account, to which all payments should be made, is as follows:",
        indent=True)

    # ── Our Client's Claim ────────────────────────────────────────────────
    _heading(doc, "Our Client's Claim", underline=True)

    _numbered_item(doc, "7.",
        "For reasons set out within this letter, it is our Client's position that you failed to "
        "comply with your statutory and regulatory obligations when considering our Client's "
        "application for finance. Specifically, you failed to carry out a proper Creditworthiness "
        "Assessment as required by the Financial Conduct Authority Consumer Credit sourcebook.",
        indent=True)

    _numbered_item(doc, "8.",
        "As a consequence of your failures, our Client is entitled to apply for remedies under CCA 1974.",
        indent=True)

    # ── Your Obligations ──────────────────────────────────────────────────
    _heading(doc, "Your Obligations", underline=True)

    _numbered_item(doc, "9.",
        "The finance that you provided to our Client is regulated by CCA 1974. Accordingly, when "
        "considering our Client's finance application, you were subject to a number of statutory "
        "and regulatory obligations.",
        indent=True)

    _numbered_item(doc, "10.",
        "CONC 5.2A.4R and 5.2A.5R require that before entering into a regulated credit agreement "
        "you must be able to demonstrate that you have undertaken a reasonable Creditworthiness "
        "Assessment and that you had reasonable grounds to be satisfied that the Credit Agreement "
        "was affordable for our Client — in particular, that our Client was in a position to make "
        "repayments as they fell due without experiencing undue difficulty.",
        indent=True)

    # Para 11: evidence-led wording — stronger for critical evidence, more proportionate otherwise
    if critical_evidence:
        para_11_text = (
            "Had you properly carried out a Creditworthiness Assessment, the significant adverse "
            "indicators already visible on our Client's credit file at the date of the Agreement "
            "would have made it apparent that the proposed credit was unaffordable and that entering "
            "into the Agreement was not in our Client's best financial interests."
        )
    else:
        para_11_text = (
            "Had you properly carried out a proportionate Creditworthiness Assessment, there were "
            "sufficient indicators of ongoing financial commitments and potential financial stress "
            "on our Client's credit file such that enhanced affordability checks ought reasonably "
            "to have been undertaken before the Agreement was entered into."
        )
    _numbered_item(doc, "11.", para_11_text, indent=True)

    _numbered_item(doc, "12.",
        "Your failure to carry out, whether properly or at all, a Creditworthiness Assessment "
        "prior to entering into the Agreement is such that it affects the fairness of the "
        "relationship between you and our Client.",
        indent=True)

    _numbered_item(doc, "13.",
        "Accordingly, our Client is entitled to apply to the Court for an order pursuant to "
        "s.140A of the Consumer Credit Act 1974 (\"CCA 1974\") that the Agreement is \"unfair\".",
        indent=True)

    _numbered_item(doc, "14.",
        "Following a finding by the Court that the Agreement is unfair, our Client will then be "
        "entitled under s140B of CCA 1974 to seek one or more of the following remedies:",
        indent=True)

    s140b_labels = ["14(a)", "14(b)", "14(c)", "14(d)", "14(e)"]
    s140b = [
        "require you to repay (in whole or in part) any sum paid by our Client by virtue of the "
        "Agreement or any related agreement (whether paid to the creditor, the associate or the "
        "former associate or to any other person) as the court thinks just; and/or",
        "require you to do or not to do (or to cease doing) anything specified in the order in "
        "connection with the Agreement or any related agreement; and/or",
        "reduce or discharge any sum payable by the debtor or by a surety by virtue of the "
        "Agreement or any related agreement; and/or",
        "otherwise set aside (in whole or in part) any duty imposed on our Client by virtue of "
        "the Agreement or any related agreement; and/or",
        "alter the terms of the Agreement or of any related agreement.",
    ]
    for lbl, item in zip(s140b_labels, s140b):
        _numbered_item(doc, lbl, item, indent=True)

    _numbered_item(doc, "15.",
        "Having taken initial instructions from our Client, it appears that you did not carry "
        "out any, or any adequate, affordability checks before entering into the Agreement with "
        "our Client.",
        indent=True)

    _numbered_item(doc, "16.",
        "It is our Client's position that you failed to comply with CONC 2.5.3R owing to the "
        "failure to carry out sufficient affordability checks, thus no reasonable steps were "
        "taken to consider whether the Agreement was suitable for our Client.",
        indent=True)

    _numbered_item(doc, "17.",
        "Pursuant to CONC 5.2A.15R(2), the firm must take reasonable steps to determine the "
        "amount, or make a reasonable estimate, of the Client's current income. We have performed "
        "an assessment of our Client's income and financial position at the time of the application "
        "and can confirm that the Agreement was not affordable.",
        indent=True)

    _numbered_item(doc, "18.",
        f"We refer to the attached affordability report (\"Report\") prepared by an independent "
        f"assessor, dated {today_str}, which outlines that the finance was Unaffordable.",
        indent=True)

    _numbered_item(doc, "19.",
        "It is our Client's position that our Client did not have an adequate discretionary income "
        "to maintain payments under the Agreement, nor to service all of their existing financial "
        "commitments.",
        indent=True)

    # Para 20 — concurrent facilities (dynamic)
    if concurrent:
        _numbered_item(doc, "20.",
            "Our Client maintained the following concurrent financial facilities at the time of the "
            "Agreement, further demonstrating the level of existing financial commitment and "
            "substantiating the unaffordability of the Agreement:",
            indent=True)
        _facilities_table(doc, concurrent)
    else:
        _numbered_item(doc, "20.",
            "As evidenced within the Report, our Client maintained concurrent credit commitments "
            "at the time of the Agreement, the nature and extent of which further substantiates "
            "the unaffordability of entering into the Agreement.",
            indent=True)

    # ── CONC 5.2A.12R obligations ─────────────────────────────────────────
    _numbered_item(doc, "21.",
        "Furthermore, CONC 5.2A.12R requires you to consider the Client's ability to make "
        "repayments as they fall due under the Agreement using the Client's income and/or income "
        "from savings and/or from savings jointly held with another person without:",
        indent=True)

    for num, text in [
        ("21.1", "having to borrow to meet the repayments;"),
        ("21.2", "failing to make any other payments that the Client is obliged to make;"),
        ("21.3", "the repayments having a significant adverse impact on the Client's financial situation."),
    ]:
        _numbered_item(doc, num, text)

    # ── Financial commitments and borrowing behaviour ─────────────────────
    _numbered_item(doc, "22.",
        "In assessing the Creditworthiness of our Client, you should have considered the following "
        "financial commitments and borrowing behaviour:",
        indent=True)

    _numbered_item(doc, "22.1", "Concurrent Loan and Credit Agreement Commitments")
    _para(doc,
        f"As shown within the Report, our Client held loans and credit agreements with other "
        f"lenders at the time of the Agreement. The accumulation of multiple credit commitments "
        f"signals that our Client was already carrying a significant financial burden at the time "
        f"of the Agreement. Such a credit profile raises serious concerns about the Lender's "
        f"failure to ensure that credit is extended in a sustainable and responsible manner.",
        size=11, space_after=6)

    _numbered_item(doc, "22.2", "Ongoing Reliance on Credit")
    _para(doc,
        "Our Client's credit file evidences a pattern of ongoing reliance on credit facilities "
        "to meet financial obligations. In the period leading up to the Agreement, our Client "
        "was consistently utilising available credit, which is consistent with a client whose "
        "income was insufficient to service their existing financial commitments without recourse "
        "to further borrowing.",
        size=11, space_after=6)

    # ── Risk parameters (dynamic, evidence-specific) ──────────────────────
    _numbered_item(doc, "23.",
        "From a credit or lending perspective, our Client's financial profile at the time of the "
        "Agreement disclosed the following risk indicators, each of which ought to have been "
        "identified and acted upon during the Creditworthiness Assessment:",
        indent=True)

    for _key, title, text in active_sections:
        _numbered_item(doc, title, "")
        _para(doc, text, size=11, space_after=6)

    # ── Lender-type specific argument paragraph ───────────────────────────
    loc_arg = get_loc_argument(lender_type, lender)
    _para(doc, loc_arg, size=11, space_after=8)

    # Paragraphs 24–28
    _numbered_item(doc, "24.",
        "The above factors should have given a clear indication to you that our Client could not "
        "reasonably afford the finance provided by you under the Agreement.",
        indent=True)

    _numbered_item(doc, "25.",
        "Following the above, it is alleged that little to no attempt was made to fairly assess "
        "our Client's creditworthiness, further solidifying our Client's position that the "
        "relationship was unfair. Moreover, any such assessment would have indicated that our "
        "Client was not in a position to service the Agreement without undue hardship.",
        indent=True)

    _numbered_item(doc, "26.",
        "In breach of CONC 5.2A.10R, it appears from the Report that you failed to consider "
        "(1) the risk that our Client would not make repayments under the Agreement by their "
        "due dates (\"credit risk\"); and (2) the impact of the Agreement on the Client's "
        "financial situation (\"financial situation risk\"). Specifically, you did not consider "
        "our Client's ability to make repayments:",
        indent=True)

    for num, text in [
        ("26.1", "as they fall due under the life of the agreement;"),
        ("26.2",
         "out of, or using, one or more of the following: (a) the Client's income; (b) income "
         "from savings or assets jointly held by the Client with another person, income received "
         "by the Client jointly with another person; (c) financial resources jointly held by the "
         "Client with another person;"),
        ("26.3", "without having to borrow to meet the repayments;"),
        ("26.4", "without failing to make any other payment the Client has a contractual or statutory obligation to make; and"),
        ("26.5", "without the repayments having a significant adverse impact on the Client's financial situation."),
    ]:
        _numbered_item(doc, num, text)

    _numbered_item(doc, "27.",
        "Had the necessary Creditworthiness Assessment been made, you would have been aware that "
        "the Agreement and the financial burden that it would create were not affordable for our "
        "Client. As a consequence, the Agreement should not have been entered into.",
        indent=True)

    _numbered_item(doc, "28.",
        "Accordingly, taking account of all of the matters set out above, we consider that the "
        "Court will find that the Agreement was unfair. We consider that the resultant award, "
        "reflecting the nature and degree of the unfairness of the relationship, will include:",
        indent=True)

    remedies_28 = [
        ("28.1",
         "repayment of the sums paid by our Client under the Agreement. It is our Client's "
         "position that had the necessary Creditworthiness Assessment been undertaken, the "
         "Agreement would never have been entered into. Our Client is therefore entitled to seek "
         "a refund of all sums paid under the Agreement;"),
        ("28.2",
         "repayment of the balance of the interest paid to you on the basis that, if proper "
         "checks had been made, then our Client might not have proceeded with the transaction;"),
        ("28.3", "where a deposit was paid, a repayment of the deposit paid under the Agreement; and"),
        ("28.4",
         "payment of simple interest, on the basis of our Client's loss of use of the money "
         "paid to you, for which the rate should be fixed by reference to interest rates for "
         "individual borrowing, given that this is what our Client would otherwise have had to "
         "pay in order to raise the funds paid to you."),
    ]
    for num, text in remedies_28:
        _numbered_item(doc, num, text)

    doc.add_paragraph()

    # ── Limitation ────────────────────────────────────────────────────────
    _heading(doc, "Limitation", underline=True)

    _numbered_item(doc, "29.",
        "The limitation period in respect of this claim is six years from the end of the "
        "relationship between you and our Client, rather than six years from the date of the "
        "Agreement, following the Supreme Court's judgment in Smith v Royal Bank of Scotland plc "
        "[2023] UKSC 34. Our Client therefore reserves their right to bring this claim within "
        "the applicable limitation period.",
        indent=True)

    # ── Disclosure ────────────────────────────────────────────────────────
    _heading(doc, "Disclosure", underline=True)

    _numbered_item(doc, "30.",
        "At this time, we are requesting from you, disclosure of all documents relating to our "
        "Client's finance application including but not limited to:",
        indent=True)

    disclosure = [
        ("30.1",  "the creditworthiness assessment that was undertaken (CONC 5.2A.4R); and"),
        ("30.2",  "the affordability risk assessment that was undertaken (CONC 5.2A.5R(2)); and"),
        ("30.3",  "all information of which you were aware at the time of the application, "
                  "including documents, supplied by the Client or from any credit reference "
                  "agency (CONC 5.2A.7); and"),
        ("30.4",  "application for finance; and"),
        ("30.5",  "details of income, including documentation, used for verification of the income; and"),
        ("30.6",  "details of expenditure, including documentation, used for verification of expenditure; and"),
        ("30.7",  "your assessment of the income and expenditure; and"),
        ("30.8",  "your documented lending criteria and parameters used to make a lending decision; and"),
        ("30.9",  "copy of the credit report used for the purpose of the Creditworthiness Assessment; and"),
        ("30.10", "copy of the statement of account; and"),
        ("30.11", "all information utilised from the Office of National Statistics to undertake "
                  "the Creditworthiness Assessment."),
    ]
    for num, text in disclosure:
        _numbered_item(doc, num, text)

    doc.add_paragraph()

    # ── Your Response ─────────────────────────────────────────────────────
    _heading(doc, "Your Response", underline=True)

    _numbered_item(doc, "31.",
        "Taking account of all of the matters raised above and the nature of your business, we "
        "consider that a reasonable time for you to respond to this Letter of Claim would be "
        "21 days from the date on which this Letter of Claim is sent to you.",
        indent=True)

    _numbered_item(doc, "32.",
        "Further, in addition to your response letter, we also require from you pre-action "
        "disclosure under CPR 31.16 of all documents setting out your creditworthiness assessment "
        "in respect of our Client, as set out in the Disclosure section above.",
        indent=True)

    _numbered_item(doc, "33.",
        "We will, of course, consider carefully the contents of your response. However, unless "
        "that response should accept our Client's claim or, alternatively, give substantive "
        "reasons for your non-acceptance of the claim, we will take such steps as may be "
        "necessary, including the commencement of proceedings.",
        indent=True)

    _numbered_item(doc, "34.",
        "Of course, our Client is mindful that litigation should be a last resort. Accordingly, "
        "our Client would be prepared to engage with you in some form of ADR to attempt "
        "settlement of the dispute without the need for Court proceedings.",
        indent=True)

    # ── Closing ───────────────────────────────────────────────────────────
    _para(doc, "Yours faithfully,", size=11, space_after=28)
    _para(doc, FIRM_NAME, bold=True, size=11, space_after=2)
    _para(doc, FIRM_ADDRESS, size=11, space_after=0)

    _add_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
