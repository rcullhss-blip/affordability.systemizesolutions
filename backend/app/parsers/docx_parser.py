import io
from docx import Document


def extract_docx(raw_bytes: bytes) -> str:
    doc = Document(io.BytesIO(raw_bytes))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)
